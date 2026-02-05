"""
Document extraction module: Extract text from PDF, DOCX, and TXT files
"""
import logging
from pathlib import Path
from typing import List, Tuple

# Try to import PDF libraries
HAS_PYPDF = False
HAS_PDFPLUMBER = False

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    pass

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    pass

# Try to import DOCX library
HAS_DOCX = False
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    pass

if not HAS_PYPDF and not HAS_PDFPLUMBER:
    raise ImportError("Please install either pypdf or pdfplumber: pip install pypdf or pip install pdfplumber")

from app.config import MAX_PAGES

logger = logging.getLogger(__name__)


def extract_pdf_text(pdf_path: Path, max_pages: int = MAX_PAGES) -> List[Tuple[int, str]]:
    """
    Extract text from PDF, only processing first max_pages pages
    
    Args:
        pdf_path: Path to PDF file
        max_pages: Maximum number of pages to process (default: 3)
    
    Returns:
        List of tuples (page_number, page_text)
    """
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")
    
    logger.info(f"Extracting text from {pdf_path}, processing first {max_pages} pages")
    
    if HAS_PYPDF:
        try:
            # Try pypdf first
            reader = PdfReader(str(pdf_path))
            pages = []
            
            for i, page in enumerate(reader.pages[:max_pages], start=1):
                try:
                    text = page.extract_text()
                    if text.strip():
                        pages.append((i, text.strip()))
                        logger.debug(f"Extracted {len(text)} characters from page {i}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {i}: {e}")
                    pages.append((i, ""))
            
            logger.info(f"Successfully extracted text from {len(pages)} pages")
            return pages
        except Exception as e:
            logger.warning(f"pypdf failed, trying pdfplumber: {e}")
    
    # Fallback to pdfplumber
    if HAS_PDFPLUMBER:
        try:
            pages = []
            
            with pdfplumber.open(str(pdf_path)) as pdf:
                for i, page in enumerate(pdf.pages[:max_pages], start=1):
                    try:
                        text = page.extract_text()
                        if text:
                            pages.append((i, text.strip()))
                            logger.debug(f"Extracted {len(text)} characters from page {i}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {i}: {e}")
                        pages.append((i, ""))
            
            logger.info(f"Successfully extracted text from {len(pages)} pages using pdfplumber")
            return pages
        except Exception as e2:
            raise RuntimeError(f"Failed to extract text from PDF using pdfplumber: {e2}")
    
    # If we get here, neither library worked
    raise RuntimeError("Failed to extract text from PDF: no working PDF library available")


def extract_full_text(pdf_path: Path, max_pages: int = MAX_PAGES) -> str:
    """
    Extract full text from PDF as a single string

    Args:
        pdf_path: Path to PDF file
        max_pages: Maximum number of pages to process

    Returns:
        Combined text from all pages
    """
    pages = extract_pdf_text(pdf_path, max_pages)
    return "\n\n".join([f"Page {page_num}:\n{text}" for page_num, text in pages])


def extract_txt_text(txt_path: Path) -> List[Tuple[int, str]]:
    """
    Extract text from plain text file

    Args:
        txt_path: Path to TXT file

    Returns:
        List with single tuple (page_number=1, full_text)
    """
    if not txt_path.exists():
        raise FileNotFoundError(f"Text file not found: {txt_path}")

    logger.info(f"Extracting text from {txt_path}")

    try:
        text = txt_path.read_text(encoding='utf-8')
        logger.info(f"Successfully extracted {len(text)} characters from text file")
        return [(1, text.strip())]
    except UnicodeDecodeError:
        # Try different encoding
        try:
            text = txt_path.read_text(encoding='latin-1')
            logger.info(f"Successfully extracted {len(text)} characters from text file (latin-1)")
            return [(1, text.strip())]
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from TXT file: {e}")


def extract_docx_text(docx_path: Path) -> List[Tuple[int, str]]:
    """
    Extract text from DOCX file

    Args:
        docx_path: Path to DOCX file

    Returns:
        List with single tuple (page_number=1, full_text)
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    if not HAS_DOCX:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")

    logger.info(f"Extracting text from {docx_path}")

    try:
        doc = Document(str(docx_path))

        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n\n".join(text_parts)
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX file")
        return [(1, full_text)]

    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOCX file: {e}")


def extract_document_text(file_path: Path, max_pages: int = MAX_PAGES) -> List[Tuple[int, str]]:
    """
    Generic document text extraction - handles PDF, TXT, and DOCX files

    Args:
        file_path: Path to document file
        max_pages: Maximum number of pages to process (for PDF only)

    Returns:
        List of tuples (page_number, page_text)
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = file_path.suffix.lower()

    if suffix == '.pdf':
        return extract_pdf_text(file_path, max_pages)
    elif suffix == '.txt':
        return extract_txt_text(file_path)
    elif suffix in ['.docx', '.doc']:
        if suffix == '.doc':
            logger.warning("DOC format detected. This requires python-docx which only handles DOCX. Please use DOCX format.")
        return extract_docx_text(file_path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Supported formats: .pdf, .txt, .docx")
